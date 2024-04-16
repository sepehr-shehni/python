import os
from docx import Document

# پیدا کردن نام کامپیوتر
computer_name = os.environ['COMPUTERNAME']

# ایجاد یک شیء Document جدید
doc = Document()

# افزودن نام کامپیوتر به سند
doc.add_heading('Computer Name', level=1)
doc.add_paragraph(computer_name)

# ذخیره سند با نام مشخص
doc.save('computer_name.docx')
