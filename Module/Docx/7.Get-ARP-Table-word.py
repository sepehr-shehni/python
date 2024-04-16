import os
from docx import Document

# پیدا کردن نام کامپیوتر
Arp_Tables = os.system("arp -a")

# ایجاد یک شیء Document جدید
doc = Document()

# افزودن نام کامپیوتر به سند
doc.add_heading('Show ARP Table', level=1)
doc.add_paragraph(Arp_Tables)

# ذخیره سند با نام مشخص
doc.save('Arp_Tables.docx')
