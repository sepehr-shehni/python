import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def create_word_file(text):
    document = Document()
    document.add_heading('Text to Word', level=1)
    document.add_paragraph(text)
    document.save('output.docx')

def convert_to_pdf():
    try:
        c = canvas.Canvas("output.pdf", pagesize=letter)
        c.setFont("Helvetica", 12)
        text = text_entry.get(1.0, tk.END)
        lines = text.split("\n")
        y = 750
        for line in lines:
            c.drawString(50, y, line)
            y -= 15
        c.save()
        messagebox.showinfo("Success", "PDF created successfully!")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

# Create the main window
root = tk.Tk()
root.title("Text to PDF")

# Text entry widget
text_entry = tk.Text(root, height=10, width=50)
text_entry.pack(pady=10)

# Button to create Word file
word_button = tk.Button(root, text="Create Word File", command=lambda: create_word_file(text_entry.get(1.0, tk.END)))
word_button.pack(pady=5)

# Button to convert Word to PDF
pdf_button = tk.Button(root, text="Convert to PDF", command=convert_to_pdf)
pdf_button.pack(pady=5)

# Run the application
root.mainloop()
